from openai import OpenAI
from fastapi import FastAPI, Form, Request, WebSocket, UploadFile
from typing import Annotated
from fastapi.templating import Jinja2Templates
from fastapi.responses import HTMLResponse
from fastapi.staticfiles import StaticFiles
import os
from dotenv import load_dotenv
from docx import Document
import fitz  # PyMuPDF to read PDF
from openpyxl import load_workbook  # For handling merged cells in Excel
import logging
import asyncio  # To handle asynchronous file processing
 
# Load environment variables from .env file
load_dotenv()
 
# Initialize OpenAI with API key from environment variable
openai = OpenAI(
    api_key=os.getenv('OPENAI_API_SECRET_KEY')
)
 
# Create FastAPI app
app = FastAPI()
 
# Mount the "templates1" directory to serve static files (like images)
app.mount("/static", StaticFiles(directory="templates1"), name="static")
 
# Specify the directory for Jinja2 templates
templates = Jinja2Templates(directory="templates1")
 
chat_responses = []
 
def load_chat_log_from_docx(file_path):
    """Load text and tables from a Word document."""
    doc = Document(file_path)
    chat_log = []
 
    # Extract text from paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            chat_log.append({'role': 'user', 'content': para.text.strip()})
 
    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                chat_log.append({'role': 'user', 'content': row_text})
 
    return chat_log
 
def load_chat_log_from_excel(file_path):
    """Load chat log by reading all worksheets from an Excel file, handling merged cells and preserving table structure."""
    chat_log = []
    try:
        # Load the workbook and access each sheet using openpyxl
        workbook = load_workbook(file_path, data_only=True)
        
        # Iterate through each sheet in the workbook
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            sheet_log = []
            
            # Read the data as a table row by row
            for row in sheet.iter_rows(values_only=False):
                row_data = []
                
                for cell in row:
                    cell_value = None
                    # If the cell is merged, get the value from the top-left cell of the merge range
                    if cell.merged_cell:
                        # Find the correct merged cell range
                        for merged_range in sheet.merged_cells.ranges:
                            if cell.coordinate in merged_range:
                                # Take the value from the top-left cell of the merged range
                                merged_cell = sheet[merged_range.min_row][merged_range.min_col - 1]
                                cell_value = merged_cell.value
                                break
                    else:
                        cell_value = cell.value
 
                    if cell_value is None:
                        cell_value = ''  # Replace None values with an empty string
                    else:
                        cell_value = str(cell_value).strip()  # Convert to string and strip extra whitespace
                    
                    row_data.append(cell_value)
 
                # Only add non-empty rows to the chat log
                if any(row_data):
                    row_text = ' | '.join(row_data)
                    sheet_log.append({'role': 'user', 'content': row_text})
 
            if sheet_log:
                chat_log.extend(sheet_log)
 
    except FileNotFoundError as fnf_error:
        logging.error(f"Excel file not found: {fnf_error}")
    except Exception as e:
        logging.error(f"Error reading Excel file {file_path}: {e}")
 
    return chat_log
 
async def load_chat_log_from_pdf(file_path):
    """Load chat log from a PDF file asynchronously."""
    chat_log = []
    try:
        pdf_document = fitz.open(file_path)
 
        for page_num in range(len(pdf_document)):
            page = pdf_document.load_page(page_num)
            text = page.get_text("text")
            if text.strip():
                chat_log.append({'role': 'user', 'content': text.strip()})
 
        # Close the document after processing
        pdf_document.close()
 
    except Exception as e:
        logging.error(f"Error reading PDF file {file_path}: {e}")
 
    return chat_log
 
async def load_chat_logs_from_folder(folder_path, batch_size=10):
    """Load chat logs by reading from all Word, Excel, and PDF files in a folder with batch processing."""
    chat_log = []
    pdf_tasks = []
 
    for filename in os.listdir(folder_path):
        file_path = os.path.join(folder_path, filename)
        extension = filename.split('.')[-1].lower()
 
        if extension == 'docx':
            chat_log.extend(load_chat_log_from_docx(file_path))
        elif extension == 'xlsx':
            chat_log.extend(load_chat_log_from_excel(file_path))
        elif extension == 'pdf':
            # For PDF files, create asynchronous tasks and process them in batches
            pdf_tasks.append(load_chat_log_from_pdf(file_path))
 
            # Process tasks in batches
            if len(pdf_tasks) == batch_size:
                results = await asyncio.gather(*pdf_tasks)
                for result in results:
                    chat_log.extend(result)
                pdf_tasks = []
 
    # Process remaining PDF tasks
    if pdf_tasks:
        results = await asyncio.gather(*pdf_tasks)
        for result in results:
            chat_log.extend(result)
 
    return chat_log
 
# Use an async startup event to ensure the function is awaited properly
@app.on_event("startup")
async def startup_event():
    global chat_log
    chat_log = await load_chat_logs_from_folder("templates1/")
 
@app.get("/", response_class=HTMLResponse)
async def chat_page(request: Request):
    """Serve the main chat page."""
    return templates.TemplateResponse("home.html", {"request": request, "chat_responses": chat_responses})
 
@app.websocket("/ws")
async def chat(websocket: WebSocket):
    """WebSocket connection for real-time chat."""
    await websocket.accept()
 
    while True:
        user_input = await websocket.receive_text()
        chat_log.append({'role': 'user', 'content': user_input})
        chat_responses.append(user_input)
 
        try:
            response = openai.chat.completions.create(
                model='gpt-3.5-turbo',
                messages=chat_log,
                temperature=0.6,
                stream=True
            )
 
            ai_response = ''
 
            for chunk in response:
                if chunk.choices[0].delta.content is not None:
                    ai_response += chunk.choices[0].delta.content
                    await websocket.send_text(chunk.choices[0].delta.content)
 
            # Append AI response to the log in sentence format
            chat_log.append({'role': 'assistant', 'content': ai_response})
            chat_responses.append(ai_response)
 
        except Exception as e:
            await websocket.send_text(f'Error: {str(e)}')
            break
 
@app.post("/", response_class=HTMLResponse)
async def chat(request: Request, user_input: Annotated[str, Form()]):
    """Handle chat interactions from the form."""
    chat_log.append({'role': 'user', 'content': user_input})
    chat_responses.append(user_input)
 
    response = openai.chat.completions.create(
        model='gpt-4',
        messages=chat_log,
        temperature=0.6
    )
 
    bot_response = response.choices[0].message.content
    chat_log.append({'role': 'assistant', 'content': bot_response})
    chat_responses.append(bot_response)
 
    # Ensure the response is a coherent sentence
    return templates.TemplateResponse("home.html", {"request": request, "chat_responses": chat_responses})
 
@app.post("/upload", response_class=HTMLResponse)
async def upload_file(request: Request, file: UploadFile):
    """Handle file uploads (Word, Excel, PDF) and load chat logs."""
    file_location = f"templates1/{file.filename}"
    with open(file_location, "wb+") as file_object:
        file_object.write(file.file.read())
 
    # Reload chat logs from all files in the templates1 folder after upload
    chat_log.extend(await load_chat_logs_from_folder("templates1/"))
 
    return templates.TemplateResponse("home.html", {"request": request, "chat_responses": chat_responses})
